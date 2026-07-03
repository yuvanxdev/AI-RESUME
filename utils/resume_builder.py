from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from io import BytesIO
import tempfile
import traceback

class ResumeBuilder:
    def __init__(self):
        self.templates = {
            "Modern": self.build_modern_template,
            "Professional": self.build_professional_template,
            "Minimal": self.build_minimal_template,
            "Creative": self.build_creative_template
        }
        
    def generate_resume(self, data):
        """Generate a resume based on the provided data and template"""
        try:
            print(f"Starting resume generation with template: {data['template']}")
            
            # Create a new document
            doc = Document()
            
            # Select and apply template
            template_name = data['template'].lower()
            print(f"Using template: {template_name}")
            
            if template_name == 'modern':
                doc = self.build_modern_template(doc, data)
            elif template_name == 'professional':
                doc = self.build_professional_template(doc, data)
            elif template_name == 'minimal':
                doc = self.build_minimal_template(doc, data)
            elif template_name == 'creative':
                doc = self.build_creative_template(doc, data)
            else:
                print(f"Warning: Unknown template '{template_name}', falling back to modern template")
                doc = self.build_modern_template(doc, data)
            
            # Save to buffer
            buffer = BytesIO()
            print("Saving document to buffer...")
            doc.save(buffer)
            buffer.seek(0)
            print("Resume generated successfully!")
            return buffer
            
        except Exception as e:
            print(f"Error in generate_resume: {str(e)}")
            print(f"Full traceback: {traceback.format_exc()}")
            print(f"Template data: {data}")
            raise

    def _format_list_items(self, items):
        """Helper function to handle both string and list inputs"""
        if isinstance(items, str):
            return [item.strip() for item in items.split('\n') if item.strip()]
        elif isinstance(items, list):
            return [item.strip() for item in items if item and item.strip()]
        return []

    def build_modern_template(self, doc, data):
        """Build modern style resume with clean, minimalist design"""
        try:
            # Set up styles
            styles = doc.styles
            
            # Name style - Modern, clean look
            name_style = styles.add_style('Modern Name', WD_STYLE_TYPE.PARAGRAPH) if 'Modern Name' not in styles else styles['Modern Name']
            name_style.font.size = Pt(24)
            name_style.font.bold = True
            name_style.font.color.rgb = RGBColor(41, 128, 185)  # Modern blue
            name_style.font.name = 'Arial'
            name_style.paragraph_format.space_after = Pt(0)
            name_style.paragraph_format.space_before = Pt(6)
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Section style - Clean and modern
            section_style = styles.add_style('Modern Section', WD_STYLE_TYPE.PARAGRAPH) if 'Modern Section' not in styles else styles['Modern Section']
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(41, 128, 185)  # Modern blue
            section_style.font.name = 'Arial'
            section_style.paragraph_format.space_before = Pt(16)
            section_style.paragraph_format.space_after = Pt(4)

            # Section underline style
            section_underline = styles.add_style('Modern Section Underline', WD_STYLE_TYPE.PARAGRAPH) if 'Modern Section Underline' not in styles else styles['Modern Section Underline']
            section_underline.font.size = Pt(8)
            section_underline.font.color.rgb = RGBColor(41, 128, 185)
            section_underline.paragraph_format.space_after = Pt(8)

            # Normal text style
            normal_style = styles.add_style('Modern Normal', WD_STYLE_TYPE.PARAGRAPH) if 'Modern Normal' not in styles else styles['Modern Normal']
            normal_style.font.size = Pt(10)
            normal_style.font.name = 'Arial'
            normal_style.paragraph_format.space_after = Pt(2)
            normal_style.font.color.rgb = RGBColor(44, 62, 80)

            # Contact style
            contact_style = styles.add_style('Modern Contact', WD_STYLE_TYPE.PARAGRAPH) if 'Modern Contact' not in styles else styles['Modern Contact']
            contact_style.font.size = Pt(10)
            contact_style.font.name = 'Arial'
            contact_style.font.color.rgb = RGBColor(41, 128, 185)
            contact_style.paragraph_format.space_after = Pt(2)
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add name at the top
            name_paragraph = doc.add_paragraph(data['personal_info']['full_name'].upper())
            name_paragraph.style = name_style

            # Add role/title if available
            if data['personal_info'].get('title'):
                title = doc.add_paragraph(data['personal_info']['title'])
                title.style = contact_style

            # Contact information layout
            contact_info = doc.add_paragraph()
            contact_info.style = contact_style
            
            # Add contact details with separators
            contact_parts = []
            if data['personal_info'].get('email'): contact_parts.append(data['personal_info']['email'])
            if data['personal_info'].get('phone'): contact_parts.append(data['personal_info']['phone'])
            if data['personal_info'].get('location'): contact_parts.append(data['personal_info']['location'])
            if contact_parts:
                contact_info.add_run(' | '.join(contact_parts))

            # Links layout
            if data['personal_info'].get('linkedin') or data['personal_info'].get('portfolio'):
                links = doc.add_paragraph()
                links.style = contact_style
                links_parts = []
                if data['personal_info'].get('linkedin'): links_parts.append(f"LinkedIn: {data['personal_info']['linkedin']}")
                if data['personal_info'].get('portfolio'): links_parts.append(f"Portfolio: {data['personal_info']['portfolio']}")
                links.add_run(' | '.join(links_parts))

            # Professional Summary
            if data.get('summary'):
                doc.add_paragraph('PROFESSIONAL SUMMARY', style=section_style)
                doc.add_paragraph('_' * 40, style=section_underline)
                summary = doc.add_paragraph(data['summary'])
                summary.style = normal_style
                summary.paragraph_format.space_after = Pt(12)
                summary.paragraph_format.left_indent = Inches(0.2)

            # Experience Section
            if data.get('experience'):
                doc.add_paragraph('EXPERIENCE', style=section_style)
                doc.add_paragraph('_' * 40, style=section_underline)
                for exp in data['experience']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.paragraph_format.left_indent = Inches(0.2)
                    
                    # Company and position
                    p.add_run(f"{exp['position']} at {exp['company']}").bold = True
                    date_run = p.add_run(f"\n{exp['start_date']} - {exp['end_date']}")
                    date_run.font.color.rgb = RGBColor(41, 128, 185)
                    
                    if exp.get('description'):
                        desc = doc.add_paragraph(exp['description'])
                        desc.style = normal_style
                        desc.paragraph_format.left_indent = Inches(0.4)
                    
                    if exp.get('responsibilities'):
                        for resp in self._format_list_items(exp['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.6)
                            bullet.add_run('• ' + resp)
                    p.paragraph_format.space_after = Pt(12)

            # Projects Section
            if data.get('projects'):
                doc.add_paragraph('PROJECTS', style=section_style)
                doc.add_paragraph('_' * 40, style=section_underline)
                for proj in data['projects']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.paragraph_format.left_indent = Inches(0.2)
                    
                    p.add_run(proj['name']).bold = True
                    if proj.get('technologies'):
                        tech_run = p.add_run(f" | {proj['technologies']}")
                        tech_run.font.color.rgb = RGBColor(41, 128, 185)
                    
                    if proj.get('description'):
                        desc = doc.add_paragraph(proj['description'])
                        desc.style = normal_style
                        desc.paragraph_format.left_indent = Inches(0.4)
                    
                    if proj.get('responsibilities'):
                        for resp in self._format_list_items(proj['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.6)
                            bullet.add_run('• ' + resp)
                    p.paragraph_format.space_after = Pt(12)

            # Education Section
            if data.get('education'):
                doc.add_paragraph('EDUCATION', style=section_style)
                doc.add_paragraph('_' * 40, style=section_underline)
                for edu in data['education']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.paragraph_format.left_indent = Inches(0.2)
                    
                    p.add_run(f"{edu['school']}").bold = True
                    p.add_run(f"\n{edu['degree']} in {edu['field']}")
                    date_run = p.add_run(f"\nGraduation: {edu['graduation_date']}")
                    if edu.get('gpa'):
                        p.add_run(f" | GPA: {edu['gpa']}")
                    p.paragraph_format.space_after = Pt(8)

            # Skills Section
            if data.get('skills'):
                doc.add_paragraph('SKILLS', style=section_style)
                doc.add_paragraph('_' * 40, style=section_underline)
                skills = data['skills']
                
                def add_skill_category(category_name, title):
                    if skills.get(category_name):
                        p = doc.add_paragraph()
                        p.style = normal_style
                        p.paragraph_format.left_indent = Inches(0.2)
                        p.add_run(f"{title}: ").bold = True
                        skills_text = ' • '.join(self._format_list_items(skills[category_name]))
                        p.add_run(skills_text)
                        p.paragraph_format.space_after = Pt(6)
                
                add_skill_category('technical', 'Technical Skills')
                add_skill_category('soft', 'Soft Skills')
                add_skill_category('languages', 'Languages')
                add_skill_category('tools', 'Tools & Technologies')

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)

            return doc
            
        except Exception as e:
            print(f"Error in build_modern_template: {str(e)}")
            raise

    def build_professional_template(self, doc, data):
        """Build professional style resume with improved spacing and layout"""
        try:
            # Set up styles
            styles = doc.styles
            
            # Header style - Name
            header_style = styles.add_style('Pro Header', WD_STYLE_TYPE.PARAGRAPH) if 'Pro Header' not in styles else styles['Pro Header']
            header_style.font.size = Pt(24)
            header_style.font.bold = True
            header_style.font.color.rgb = RGBColor(0, 0, 0)
            header_style.paragraph_format.space_after = Pt(4)
            header_style.font.name = 'Calibri'

            # Section style
            section_style = styles.add_style('Pro Section', WD_STYLE_TYPE.PARAGRAPH) if 'Pro Section' not in styles else styles['Pro Section']
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(0, 120, 215)
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
            section_style.font.name = 'Calibri'

            # Normal text style
            normal_style = styles.add_style('Pro Normal', WD_STYLE_TYPE.PARAGRAPH) if 'Pro Normal' not in styles else styles['Pro Normal']
            normal_style.font.size = Pt(10)
            normal_style.font.name = 'Calibri'
            normal_style.paragraph_format.space_after = Pt(2)

            # Contact style
            contact_style = styles.add_style('Pro Contact', WD_STYLE_TYPE.PARAGRAPH) if 'Pro Contact' not in styles else styles['Pro Contact']
            contact_style.font.size = Pt(10)
            contact_style.font.name = 'Calibri'
            contact_style.paragraph_format.space_after = Pt(6)

            # Add name at the top
            name_paragraph = doc.add_paragraph(data['personal_info']['full_name'])
            name_paragraph.style = header_style
            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Add contact information in a single line
            contact_parts = []
            if data['personal_info'].get('email'): contact_parts.append(data['personal_info']['email'])
            if data['personal_info'].get('phone'): contact_parts.append(data['personal_info']['phone'])
            if data['personal_info'].get('location'): contact_parts.append(data['personal_info']['location'])
            
            if contact_parts:
                contact = doc.add_paragraph()
                contact.style = contact_style
                contact.add_run(' | '.join(contact_parts))

            # Add LinkedIn and Portfolio links
            links_parts = []
            if data['personal_info'].get('linkedin'): links_parts.append(f"LinkedIn: {data['personal_info']['linkedin']}")
            if data['personal_info'].get('portfolio'): links_parts.append(f"Portfolio: {data['personal_info']['portfolio']}")
            
            if links_parts:
                links = doc.add_paragraph()
                links.style = contact_style
                links.add_run(' | '.join(links_parts))

            # Professional Summary
            if data.get('summary'):
                doc.add_paragraph('PROFESSIONAL SUMMARY', style=section_style)
                summary = doc.add_paragraph(data['summary'])
                summary.style = normal_style

            # Experience Section
            if data.get('experience'):
                doc.add_paragraph('EXPERIENCE', style=section_style)
                for exp in data['experience']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.add_run(f"{exp['position']} at {exp['company']}").bold = True
                    p.add_run(f" | {exp['start_date']} - {exp['end_date']}")
                    
                    if exp.get('description'):
                        desc = doc.add_paragraph(exp['description'])
                        desc.style = normal_style
                        desc.paragraph_format.left_indent = Inches(0.2)
                    
                    if exp.get('responsibilities'):
                        for resp in self._format_list_items(exp['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.3)
                            bullet.add_run('• ' + resp)

            # Projects Section
            if data.get('projects'):
                doc.add_paragraph('PROJECTS', style=section_style)
                for proj in data['projects']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.add_run(proj['name']).bold = True
                    if proj.get('technologies'):
                        p.add_run(f" | {proj['technologies']}")
                    
                    if proj.get('description'):
                        desc = doc.add_paragraph(proj['description'])
                        desc.style = normal_style
                        desc.paragraph_format.left_indent = Inches(0.2)
                    
                    if proj.get('responsibilities'):
                        for resp in self._format_list_items(proj['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.3)
                            bullet.add_run('• ' + resp)

            # Education Section
            if data.get('education'):
                doc.add_paragraph('EDUCATION', style=section_style)
                for edu in data['education']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.add_run(f"{edu['school']}").bold = True
                    p.add_run(f"\n{edu['degree']} in {edu['field']}")
                    p.add_run(f" | Graduation: {edu['graduation_date']}")
                    if edu.get('gpa'):
                        p.add_run(f" | GPA: {edu['gpa']}")

            # Skills Section
            if data.get('skills'):
                doc.add_paragraph('SKILLS', style=section_style)
                skills = data['skills']
                
                def add_skill_category(category_name, title):
                    if skills.get(category_name):
                        p = doc.add_paragraph()
                        p.style = normal_style
                        p.add_run(f"{title}: ").bold = True
                        skills_text = ', '.join(self._format_list_items(skills[category_name]))
                        p.add_run(skills_text)
                
                add_skill_category('technical', 'Technical Skills')
                add_skill_category('soft', 'Soft Skills')
                add_skill_category('languages', 'Languages')
                add_skill_category('tools', 'Tools & Technologies')

            # Set margins for better space utilization
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)

            return doc
            
        except Exception as e:
            print(f"Error in build_professional_template: {str(e)}")
            raise

    def build_minimal_template(self, doc, data):
        """Build minimal style resume"""
        try:
            # Set up styles
            styles = doc.styles
            
            # Header style - Large, bold name
            header_style = None
            if 'Min Header' not in styles:
                header_style = styles.add_style('Min Header', WD_STYLE_TYPE.PARAGRAPH)
                header_style.font.size = Pt(28)
                header_style.font.bold = True
                header_style.font.color.rgb = RGBColor(33, 33, 33)  # Dark gray
                header_style.paragraph_format.space_after = Pt(4)
            else:
                header_style = styles['Min Header']
            
            # Contact style - Small, gray text
            contact_style = None
            if 'Min Contact' not in styles:
                contact_style = styles.add_style('Min Contact', WD_STYLE_TYPE.PARAGRAPH)
                contact_style.font.size = Pt(9)
                contact_style.font.color.rgb = RGBColor(100, 100, 100)  # Light gray
                contact_style.paragraph_format.space_after = Pt(12)
            else:
                contact_style = styles['Min Contact']
            
            # Section style - Medium, all caps
            section_style = None
            if 'Min Section' not in styles:
                section_style = styles.add_style('Min Section', WD_STYLE_TYPE.PARAGRAPH)
                section_style.font.size = Pt(12)
                section_style.font.all_caps = True
                section_style.font.bold = True
                section_style.font.color.rgb = RGBColor(33, 33, 33)
                section_style.paragraph_format.space_before = Pt(16)
                section_style.paragraph_format.space_after = Pt(8)
            else:
                section_style = styles['Min Section']
            
            # Normal text style
            normal_style = None
            if 'Min Normal' not in styles:
                normal_style = styles.add_style('Min Normal', WD_STYLE_TYPE.PARAGRAPH)
                normal_style.font.size = Pt(10)
                normal_style.font.color.rgb = RGBColor(33, 33, 33)
                normal_style.paragraph_format.space_after = Pt(4)
            else:
                normal_style = styles['Min Normal']
            
            # Add header with personal info
            personal = data['personal_info']
            name = doc.add_paragraph(personal['full_name'])
            name.style = header_style
            
            # Contact info in one line
            contact_parts = []
            if personal.get('email'): contact_parts.append(personal['email'])
            if personal.get('phone'): contact_parts.append(personal['phone'])
            if personal.get('location'): contact_parts.append(personal['location'])
            
            if contact_parts:
                contact = doc.add_paragraph()
                contact.style = contact_style
                contact.add_run(' • '.join(contact_parts))
            
            # Links in one line
            links_parts = []
            if personal.get('linkedin'): links_parts.append(f"LinkedIn: {personal['linkedin']}")
            if personal.get('portfolio'): links_parts.append(f"Portfolio: {personal['portfolio']}")
            
            if links_parts:
                links = doc.add_paragraph()
                links.style = contact_style
                links.add_run(' • '.join(links_parts))
            
            # Professional Summary
            if data.get('summary'):
                doc.add_paragraph('SUMMARY', style=section_style)
                summary = doc.add_paragraph(data['summary'])
                summary.style = normal_style
            
            # Experience Section
            if data.get('experience'):
                doc.add_paragraph('EXPERIENCE', style=section_style)
                for exp in data['experience']:
                    p = doc.add_paragraph(style=normal_style)
                    p.add_run(f"{exp['position']} at {exp['company']}").bold = True
                    p.add_run(f"\n{exp['start_date']} - {exp['end_date']}")
                    
                    if exp.get('description'):  # Changed from 'overview' to 'description'
                        overview = doc.add_paragraph(exp['description'])
                        overview.style = normal_style
                    
                    if exp.get('responsibilities'):
                        resp_para = doc.add_paragraph(style=normal_style)
                        resp_para.add_run('Key Responsibilities:').bold = True
                        for resp in self._format_list_items(exp['responsibilities']):
                            bullet = doc.add_paragraph(style=normal_style)
                            bullet.style.paragraph_format.left_indent = Inches(0.25)
                            bullet.add_run('• ' + resp)
                    
                    if exp.get('achievements'):
                        ach_para = doc.add_paragraph(style=normal_style)
                        ach_para.add_run('Key Achievements:').bold = True
                        for ach in self._format_list_items(exp['achievements']):
                            bullet = doc.add_paragraph(style=normal_style)
                            bullet.style.paragraph_format.left_indent = Inches(0.25)
                            bullet.add_run('• ' + ach)
            
            # Projects Section
            if data.get('projects'):
                doc.add_paragraph('PROJECTS', style=section_style)
                for proj in data['projects']:
                    p = doc.add_paragraph(style=normal_style)
                    p.add_run(proj['name']).bold = True
                    if proj.get('technologies'):
                        p.add_run(f"\nTechnologies: {proj['technologies']}")
                    
                    if proj.get('description'):  # Changed from 'overview' to 'description'
                        overview = doc.add_paragraph(proj['description'])
                        overview.style = normal_style
                    
                    if proj.get('responsibilities'):
                        resp_para = doc.add_paragraph(style=normal_style)
                        resp_para.add_run('Key Responsibilities:').bold = True
                        for resp in self._format_list_items(proj['responsibilities']):
                            bullet = doc.add_paragraph(style=normal_style)
                            bullet.style.paragraph_format.left_indent = Inches(0.25)
                            bullet.add_run('• ' + resp)
                    
                    if proj.get('achievements'):
                        ach_para = doc.add_paragraph(style=normal_style)
                        ach_para.add_run('Key Achievements:').bold = True
                        for ach in self._format_list_items(proj['achievements']):
                            bullet = doc.add_paragraph(style=normal_style)
                            bullet.style.paragraph_format.left_indent = Inches(0.25)
                            bullet.add_run('• ' + ach)
                    
                    if proj.get('link'):
                        link = doc.add_paragraph(f"Project Link: {proj['link']}")
                        link.style = normal_style
            
            # Education Section
            if data.get('education'):
                doc.add_paragraph('EDUCATION', style=section_style)
                for edu in data['education']:
                    p = doc.add_paragraph(style=normal_style)
                    p.add_run(f"{edu['school']} - {edu['degree']} in {edu['field']}").bold = True
                    p.add_run(f"\nGraduation: {edu['graduation_date']}")
                    if edu.get('gpa'):
                        p.add_run(f" | GPA: {edu['gpa']}")
                    
                    if edu.get('achievements'):
                        ach_para = doc.add_paragraph(style=normal_style)
                        ach_para.add_run('Achievements & Activities:').bold = True
                        for ach in self._format_list_items(edu['achievements']):
                            bullet = doc.add_paragraph(style=normal_style)
                            bullet.style.paragraph_format.left_indent = Inches(0.25)
                            bullet.add_run('• ' + ach)
            
            # Skills Section
            if data.get('skills'):
                doc.add_paragraph('SKILLS', style=section_style)
                skills = data['skills']
                
                def add_skill_category(category_name, title):
                    if skills.get(category_name):
                        p = doc.add_paragraph(style=normal_style)
                        p.add_run(f"{title}: ").bold = True
                        p.add_run(' • '.join(self._format_list_items(skills[category_name])))
                
                add_skill_category('technical', 'Technical Skills')
                add_skill_category('soft', 'Soft Skills')
                add_skill_category('languages', 'Languages')
                add_skill_category('tools', 'Tools & Technologies')
            
            return doc
            
        except Exception as e:
            print(f"Error in build_minimal_template: {str(e)}")
            raise

    def build_creative_template(self, doc, data):
        """Build creative style resume with vibrant design and emojis"""
        try:
            # Set up styles
            styles = doc.styles
            
            # Name style - Creative and bold
            name_style = styles.add_style('Creative Name', WD_STYLE_TYPE.PARAGRAPH) if 'Creative Name' not in styles else styles['Creative Name']
            name_style.font.size = Pt(24)
            name_style.font.bold = True
            name_style.font.color.rgb = RGBColor(155, 89, 182)  # Purple
            name_style.font.name = 'Arial'
            name_style.paragraph_format.space_after = Pt(4)
            name_style.paragraph_format.space_before = Pt(6)
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Section style - Vibrant
            section_style = styles.add_style('Creative Section', WD_STYLE_TYPE.PARAGRAPH) if 'Creative Section' not in styles else styles['Creative Section']
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(155, 89, 182)  # Purple
            section_style.font.name = 'Arial'
            section_style.paragraph_format.space_before = Pt(16)
            section_style.paragraph_format.space_after = Pt(4)

            # Normal text style - Clean
            normal_style = styles.add_style('Creative Normal', WD_STYLE_TYPE.PARAGRAPH) if 'Creative Normal' not in styles else styles['Creative Normal']
            normal_style.font.size = Pt(10)
            normal_style.font.name = 'Arial'
            normal_style.paragraph_format.space_after = Pt(2)
            normal_style.font.color.rgb = RGBColor(52, 73, 94)  # Dark slate

            # Contact style - Professional
            contact_style = styles.add_style('Creative Contact', WD_STYLE_TYPE.PARAGRAPH) if 'Creative Contact' not in styles else styles['Creative Contact']
            contact_style.font.size = Pt(10)
            contact_style.font.name = 'Arial'
            contact_style.font.color.rgb = RGBColor(155, 89, 182)  # Purple
            contact_style.paragraph_format.space_after = Pt(2)
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add name at the top
            name_paragraph = doc.add_paragraph('✨ ' + data['personal_info']['full_name'] + ' ✨')
            name_paragraph.style = name_style

            # Add role/title if available
            if data['personal_info'].get('title'):
                title = doc.add_paragraph('💫 ' + data['personal_info']['title'])
                title.style = contact_style

            # Contact information layout
            contact_info = doc.add_paragraph()
            contact_info.style = contact_style
            
            contact_parts = []
            if data['personal_info'].get('email'): contact_parts.append(f"📧 {data['personal_info']['email']}")
            if data['personal_info'].get('phone'): contact_parts.append(f"📱 {data['personal_info']['phone']}")
            if data['personal_info'].get('location'): contact_parts.append(f"📍 {data['personal_info']['location']}")
            if contact_parts:
                contact_info.add_run(' | '.join(contact_parts))

            # Links with professional formatting
            if data['personal_info'].get('linkedin') or data['personal_info'].get('portfolio'):
                links = doc.add_paragraph()
                links.style = contact_style
                links_parts = []
                if data['personal_info'].get('linkedin'): links_parts.append(f"🔗 LinkedIn: {data['personal_info']['linkedin']}")
                if data['personal_info'].get('portfolio'): links_parts.append(f"🌐 Portfolio: {data['personal_info']['portfolio']}")
                links.add_run(' | '.join(links_parts))

            # Professional Summary
            if data.get('summary'):
                doc.add_paragraph('👨‍💼 PROFESSIONAL SUMMARY', style=section_style)
                summary = doc.add_paragraph(data['summary'])
                summary.style = normal_style
                summary.paragraph_format.space_after = Pt(12)
                summary.paragraph_format.left_indent = Inches(0.2)

            # Experience Section
            if data.get('experience'):
                doc.add_paragraph('💼 EXPERIENCE', style=section_style)
                for exp in data['experience']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.paragraph_format.left_indent = Inches(0.2)
                    
                    p.add_run(f"🚀 {exp['position']}").bold = True
                    p.add_run(f"\n🏢 {exp['company']}")
                    p.add_run(f"\n📅 {exp['start_date']} - {exp['end_date']}")
                    
                    if exp.get('description'):
                        desc = doc.add_paragraph(exp['description'])
                        desc.style = normal_style
                        desc.paragraph_format.left_indent = Inches(0.4)
                    
                    if exp.get('responsibilities'):
                        resp_para = doc.add_paragraph()
                        resp_para.style = normal_style
                        resp_para.paragraph_format.left_indent = Inches(0.4)
                        resp_para.add_run('🎯 Key Achievements:').bold = True
                        for resp in self._format_list_items(exp['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.6)
                            bullet.add_run('• ' + resp)
                    p.paragraph_format.space_after = Pt(12)

            # Projects Section
            if data.get('projects'):
                doc.add_paragraph('🛠️ PROJECTS', style=section_style)
                for proj in data['projects']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.paragraph_format.left_indent = Inches(0.2)
                    
                    p.add_run(f"✨ {proj['name']}").bold = True
                    if proj.get('technologies'):
                        p.add_run(f"\n💻 Technologies: {proj['technologies']}")
                    
                    if proj.get('description'):
                        desc = doc.add_paragraph(proj['description'])
                        desc.style = normal_style
                        desc.paragraph_format.left_indent = Inches(0.4)
                    
                    if proj.get('responsibilities'):
                        resp_para = doc.add_paragraph()
                        resp_para.style = normal_style
                        resp_para.paragraph_format.left_indent = Inches(0.4)
                        resp_para.add_run('🎯 Key Features:').bold = True
                        for resp in self._format_list_items(proj['responsibilities']):
                            bullet = doc.add_paragraph()
                            bullet.style = normal_style
                            bullet.paragraph_format.left_indent = Inches(0.6)
                            bullet.add_run('• ' + resp)
                    p.paragraph_format.space_after = Pt(12)

            # Education Section
            if data.get('education'):
                doc.add_paragraph('🎓 EDUCATION', style=section_style)
                for edu in data['education']:
                    p = doc.add_paragraph()
                    p.style = normal_style
                    p.paragraph_format.left_indent = Inches(0.2)
                    
                    p.add_run(f"📚 {edu['school']}").bold = True
                    p.add_run(f"\n🎯 {edu['degree']} in {edu['field']}")
                    p.add_run(f"\n📅 Graduation: {edu['graduation_date']}")
                    if edu.get('gpa'):
                        p.add_run(f" | 📊 GPA: {edu['gpa']}")
                    p.paragraph_format.space_after = Pt(8)

            # Skills Section
            if data.get('skills'):
                doc.add_paragraph('⭐ SKILLS', style=section_style)
                skills = data['skills']
                
                def add_skill_category(category_name, title, icon):
                    if skills.get(category_name):
                        p = doc.add_paragraph()
                        p.style = normal_style
                        p.paragraph_format.left_indent = Inches(0.2)
                        p.add_run(f"{icon} {title}: ").bold = True
                        skills_text = ' • '.join(self._format_list_items(skills[category_name]))
                        p.add_run(skills_text)
                        p.paragraph_format.space_after = Pt(6)
                
                add_skill_category('technical', 'Technical Skills', '💻')
                add_skill_category('soft', 'Soft Skills', '🤝')
                add_skill_category('languages', 'Languages', '🌐')
                add_skill_category('tools', 'Tools & Technologies', '🛠️')

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)

            return doc
            
        except Exception as e:
            print(f"Error in build_creative_template: {str(e)}")
            raise

    def generate_preview(self, template_name, data):
        """Generate a live preview of the resume"""
        if template_name not in self.preview_templates:
            return None
            
        template = self.preview_templates[template_name]
        
        # Format skills as HTML
        skills_html = ""
        if 'skills' in data:
            if template_name == 'Modern':
                skills_html = "".join([f'<div class="skill">{skill}</div>' for skill in data['skills']])
            else:
                skills_html = "".join([f'<div class="skill-item">{skill}</div>' for skill in data['skills']])
        
        # Format experience as HTML
        experience_html = ""
        if 'experience' in data:
            for exp in data['experience']:
                experience_html += f"""
                <div class="experience-item">
                    <h3>{exp.get('title', '')}</h3>
                    <p class="company">{exp.get('company', '')}</p>
                    <p class="date">{exp.get('date', '')}</p>
                    <p class="description">{exp.get('description', '')}</p>
                </div>
                """
        
        # Format education as HTML
        education_html = ""
        if 'education' in data:
            for edu in data['education']:
                education_html += f"""
                <div class="education-item">
                    <h3>{edu.get('degree', '')}</h3>
                    <p class="school">{edu.get('school', '')}</p>
                    <p class="date">{edu.get('date', '')}</p>
                </div>
                """
        
        # Combine HTML and CSS
        preview_html = template['html'].format(
            name=data.get('name', 'Your Name'),
            email=data.get('email', 'email@example.com'),
            phone=data.get('phone', '123-456-7890'),
            linkedin=data.get('linkedin', 'linkedin.com/in/yourprofile'),
            title=data.get('title', 'Your Title'),
            summary=data.get('summary', 'Your professional summary...'),
            experience=experience_html,
            education=education_html,
            skills=skills_html
        )
        
        return {
            'html': preview_html,
            'css': template['css']
        }
